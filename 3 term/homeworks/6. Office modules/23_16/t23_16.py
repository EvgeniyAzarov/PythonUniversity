import openpyxl
from docx import Document
from collections import namedtuple


Shipper = namedtuple("Shipper", "sid name rating address")


class Product():
    A1 = 0.5
    A2 = 0.5
    
    def __init__(self, pid, name):
        self.pid = pid 
        self.name = name 
        
        self._offers = {}
        self._max_price = 0
        self._max_rating = 0
    
    def add_offer(self, sid, price, term, rating):
        Offer = namedtuple("Offer", "price term rating")
        self._offers[sid] = Offer(price, term, rating)
        
        if price > self._max_price:
            self._max_price = price
            
        if rating > self._max_rating:
            self._max_rating = rating

    def _offer_priority(self, offer):
        return (Product.A1 * offer.price / self._max_price +
                Product.A2 * offer.rating / self._max_rating)

    def get_offers_top_sids(self, count):
        offers = []
        for sid, offer in self._offers.items():
            priority = self._offer_priority(offer)
            offers.append((sid, priority))

        sids = sorted(offers, key=lambda item: item[0], reverse=True) 
        sids = sids[:count]
        return [sid for sid, priority in sids]


def read_data(data):
    wb = openpyxl.load_workbook(tender_data) 

    shippers = {}
    sheet = wb['shipper']
    for row in sheet.rows:
        if row[0].value == "id":
            continue
        sid, name, rating, address = [_.value for _ in row]
        rating = float(rating.replace(',', '.'))
        shippers[sid] = Shipper(sid, name, rating, address)

    products = {}
    sheet = wb['product']
    for row in sheet.rows:
        if row[0].value == "id":
            continue
        pid, name = [_.value for _ in row]
        products[pid] = Product(pid, name)

    sheet = wb['price']
    for row in sheet.rows:
        if row[0].value == "S_id":
            continue
        sid, pid, price, term = [_.value for _ in row]
        price = float(price.replace(',', '.'))
        term = float(term.replace(',', '.'))
        shipper_rating = shippers[sid].rating
        products[pid].add_offer(sid, price, term, shipper_rating) 

    return shippers, products
 

if __name__ == '__main__':
    tender_data = "test_data/tender.xlsx"

    shippers, products = read_data(tender_data)
    letters = {}

    for product in products.values():
        sids = product.get_offers_top_sids(count=3)
        for sid in sids:
            if sid in letters.keys():
                letters[sid].append(product.name)
            else:
                letters[sid] = [product.name]

    for sid, products_names in letters.items():
        doc = Document("test_data/template.docx")
        letter_name = "to_" + shippers[sid].address + ".docx"

        for name in products_names:
            p = doc.add_paragraph(style='List')
            r = p.add_run()
            r.add_text(name + "\n")

        doc.save("letters/" + letter_name)
   
